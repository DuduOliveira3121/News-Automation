"""Testes unitários do NewsParser (detecção por heurísticas).

Os testes de extração em memória constroem documentos Word via io.BytesIO
para garantir isolamento total do sistema de arquivos.
Os testes de integração usam o arquivo real
``data/uploads/BOLETIM DIA 14 DE JULHO de 2026 – TERÇA-FEIRA.docx``.
"""
from __future__ import annotations

import io
import logging
from pathlib import Path
from unittest.mock import patch

import docx
import pytest

from src.infrastructure.parsers.news_parser import (
    NewsParser,
    Noticia,
    _is_all_upper,
    _is_bloco_marker,
    _is_title_candidate,
    _table_to_text,
)

# ---------------------------------------------------------------------------
# Constantes de apoio
# ---------------------------------------------------------------------------

#: Caminho para o boletim real usado nos testes de integração.
BOLETIM_PATH = Path(
    "data/uploads/BOLETIM DIA 14 DE JULHO de 2026 – TERÇA-FEIRA.docx"
)

#: Títulos em caixa alta com >= 5 palavras para usar nos documentos em memória.
TITLE_A = "PRIMEIRA NOTÍCIA SOBRE ASSUNTO IMPORTANTE DESTAQUE"
TITLE_B = "SEGUNDA NOTÍCIA COM TEMA MUITO RELEVANTE"
TITLE_C = "TERCEIRA NOTÍCIA TRATA DE CONTEÚDO ESPECÍFICO"


# ---------------------------------------------------------------------------
# Helpers de fixture
# ---------------------------------------------------------------------------


def _build_docx(*items: tuple[str, list[str]]) -> io.BytesIO:
    """Cria um .docx em memória a partir de pares (título, [parágrafos]).

    Como o novo parser não usa estilos, os títulos são inseridos como
    parágrafos normais.  Para serem detectados como títulos, devem estar
    em caixa alta e ter >= 5 palavras.
    """
    document = docx.Document()
    for title, body_paragraphs in items:
        document.add_paragraph(title)
        for para in body_paragraphs:
            document.add_paragraph(para)
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


def _parser_from_buf(buf: io.BytesIO, filename: str = "test.docx") -> list[Noticia]:
    """Abre um buffer BytesIO como Document e executa o parser."""
    document = docx.Document(buf)
    parser = NewsParser()
    return parser._extract(document, source_name=filename)


# ---------------------------------------------------------------------------
# Testes da dataclass Noticia
# ---------------------------------------------------------------------------


class TestNoticia:
    """Testes básicos da dataclass Noticia."""

    def test_fields_are_assigned_correctly(self) -> None:
        n = Noticia(
            titulo="T",
            corpo="C",
            arquivo_origem="f.docx",
            indice=1,
            posicao_inicial=5,
            posicao_final=10,
        )

        assert n.titulo == "T"
        assert n.corpo == "C"
        assert n.arquivo_origem == "f.docx"
        assert n.indice == 1
        assert n.posicao_inicial == 5
        assert n.posicao_final == 10

    def test_default_indice_is_zero(self) -> None:
        n = Noticia(titulo="T", corpo="C", arquivo_origem="f.docx")

        assert n.indice == 0

    def test_default_positions_are_zero(self) -> None:
        n = Noticia(titulo="T", corpo="C", arquivo_origem="f.docx")

        assert n.posicao_inicial == 0
        assert n.posicao_final == 0


# ---------------------------------------------------------------------------
# Testes das funções heurísticas
# ---------------------------------------------------------------------------


class TestIsBloco:
    """Testes para _is_bloco_marker."""

    def test_primeiro_bloco(self) -> None:
        assert _is_bloco_marker("1º BLOCO – EVOLUÇÃO") is True

    def test_segundo_bloco_com_espacos(self) -> None:
        assert _is_bloco_marker("2º BLOCO -  FASTBRAILLE") is True

    def test_terceiro_bloco(self) -> None:
        assert _is_bloco_marker("3º BLOCO -  TROPICAL FM") is True

    def test_bloco_com_ordinal_alternativo(self) -> None:
        assert _is_bloco_marker("4° BLOCO") is True

    def test_texto_comum_nao_e_bloco(self) -> None:
        assert _is_bloco_marker("NOTÍCIA QUALQUER") is False

    def test_bloco_no_meio_da_frase_nao_e_marcador(self) -> None:
        assert _is_bloco_marker("Texto 1º BLOCO no meio") is False

    def test_string_vazia_nao_e_bloco(self) -> None:
        assert _is_bloco_marker("") is False


class TestIsAllUpper:
    """Testes para _is_all_upper."""

    def test_texto_completamente_maiusculo(self) -> None:
        assert _is_all_upper("BANCO") is True

    def test_texto_com_acento_maiusculo(self) -> None:
        assert _is_all_upper("REGULAMENTAÇÃO") is True

    def test_texto_com_minusculas(self) -> None:
        assert _is_all_upper("Banco") is False

    def test_texto_todo_minusculo(self) -> None:
        assert _is_all_upper("banco") is False

    def test_texto_sem_letras_retorna_false(self) -> None:
        assert _is_all_upper("123 !@#") is False

    def test_string_vazia_retorna_false(self) -> None:
        assert _is_all_upper("") is False

    def test_pontuacao_e_numeros_ignorados(self) -> None:
        assert _is_all_upper("R$ 80 MIL") is True


class TestIsTitleCandidate:
    """Testes para _is_title_candidate."""

    def test_titulo_valido_retorna_true(self) -> None:
        assert _is_title_candidate(TITLE_A) is True

    def test_titulo_com_cinco_palavras_exatas(self) -> None:
        assert _is_title_candidate("CINCO PALAVRAS EM CAIXA ALTA") is True

    def test_titulo_com_menos_de_cinco_palavras_retorna_false(self) -> None:
        assert _is_title_candidate("APENAS QUATRO PALAVRAS AQUI") is False

    def test_titulo_com_minusculas_retorna_false(self) -> None:
        assert _is_title_candidate("Título com minúsculas válido") is False

    def test_string_vazia_retorna_false(self) -> None:
        assert _is_title_candidate("") is False

    def test_marcador_de_bloco_retorna_false(self) -> None:
        assert _is_title_candidate("1º BLOCO – EVOLUÇÃO") is False

    def test_tres_palavras_maiusculas_retorna_false(self) -> None:
        assert _is_title_candidate("BIOMOB SEÇÃO CURTA") is False

    def test_abner_instrucao_retorna_false(self) -> None:
        # "ABNER MOSTRAR O VÍDEO" tem apenas 4 palavras – não é título.
        assert _is_title_candidate("ABNER MOSTRAR O VÍDEO") is False


# ---------------------------------------------------------------------------
# Testes do NewsParser — validação de arquivos
# ---------------------------------------------------------------------------


class TestNewsParserValidation:
    """Testes de validação de entrada do NewsParser."""

    def test_raises_file_not_found_for_missing_file(self, tmp_path: Path) -> None:
        parser = NewsParser()
        missing = tmp_path / "nao_existe.docx"

        with pytest.raises(FileNotFoundError, match="Arquivo não encontrado"):
            parser.parse(missing)

    def test_raises_value_error_for_wrong_extension(self, tmp_path: Path) -> None:
        parser = NewsParser()
        wrong = tmp_path / "arquivo.txt"
        wrong.touch()

        with pytest.raises(ValueError, match="Extensão inválida"):
            parser.parse(wrong)

    def test_raises_value_error_for_corrupted_file(self, tmp_path: Path) -> None:
        parser = NewsParser()
        corrupted = tmp_path / "corrompido.docx"
        corrupted.write_bytes(b"not a valid docx content")

        with pytest.raises(ValueError, match="Não foi possível abrir"):
            parser.parse(corrupted)


# ---------------------------------------------------------------------------
# Testes do NewsParser — extração de conteúdo (documentos em memória)
# ---------------------------------------------------------------------------


class TestNewsParserExtraction:
    """Testes de extração com documentos criados em memória."""

    def test_single_news_is_extracted(self) -> None:
        buf = _build_docx((TITLE_A, ["Parágrafo 1.", "Parágrafo 2."]))
        noticias = _parser_from_buf(buf)

        assert len(noticias) == 1
        assert noticias[0].titulo == TITLE_A

    def test_multiple_news_are_extracted_in_order(self) -> None:
        buf = _build_docx(
            (TITLE_A, ["Corpo A"]),
            (TITLE_B, ["Corpo B"]),
            (TITLE_C, ["Corpo C"]),
        )
        noticias = _parser_from_buf(buf)

        assert len(noticias) == 3
        assert [n.titulo for n in noticias] == [TITLE_A, TITLE_B, TITLE_C]

    def test_indice_is_sequential_starting_at_one(self) -> None:
        buf = _build_docx(
            (TITLE_A, ["Corpo"]),
            (TITLE_B, ["Corpo"]),
        )
        noticias = _parser_from_buf(buf)

        assert [n.indice for n in noticias] == [1, 2]

    def test_body_paragraphs_are_joined_with_double_newline(self) -> None:
        buf = _build_docx((TITLE_A, ["Primeiro parágrafo.", "Segundo parágrafo."]))
        noticias = _parser_from_buf(buf)

        assert noticias[0].corpo == "Primeiro parágrafo.\n\nSegundo parágrafo."

    def test_empty_body_paragraphs_are_ignored(self) -> None:
        buf = _build_docx((TITLE_A, ["", "Parágrafo real.", ""]))
        noticias = _parser_from_buf(buf)

        assert noticias[0].corpo == "Parágrafo real."

    def test_news_with_no_body_has_empty_corpo(self) -> None:
        buf = _build_docx((TITLE_A, []))
        noticias = _parser_from_buf(buf)

        assert len(noticias) == 1
        assert noticias[0].corpo == ""

    def test_document_with_no_titles_returns_empty_list(self) -> None:
        document = docx.Document()
        document.add_paragraph("Apenas texto comum sem título em caixa alta.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert noticias == []

    def test_source_file_is_preserved_in_noticia(self) -> None:
        buf = _build_docx((TITLE_A, ["Corpo"]))
        noticias = _parser_from_buf(buf, filename="boletim.docx")

        assert noticias[0].arquivo_origem == "boletim.docx"

    def test_short_upper_paragraph_is_not_a_title(self) -> None:
        """Menos de 5 palavras em caixa alta → não é título, é corpo."""
        document = docx.Document()
        document.add_paragraph(TITLE_A)
        document.add_paragraph("SEÇÃO CURTA")  # 2 palavras – não é título
        document.add_paragraph("Parágrafo do corpo.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert len(noticias) == 1
        assert "SEÇÃO CURTA" in noticias[0].corpo

    def test_mixed_case_paragraph_is_body_not_title(self) -> None:
        document = docx.Document()
        document.add_paragraph(TITLE_A)
        document.add_paragraph("Parágrafo em caixa mista não é título.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert len(noticias) == 1
        assert "Parágrafo em caixa mista" in noticias[0].corpo

    def test_para_before_first_title_is_ignored_when_no_bloco(self) -> None:
        """Sem marcador de bloco, parágrafos antes do primeiro título são ignorados."""
        document = docx.Document()
        document.add_paragraph("Texto sem título antes.")
        document.add_paragraph(TITLE_A)
        document.add_paragraph("Corpo da notícia.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert len(noticias) == 1
        assert "Texto sem título antes." not in noticias[0].corpo

    def test_positions_are_set_correctly(self) -> None:
        """posicao_inicial e posicao_final são índices de bloco corretos."""
        buf = _build_docx(
            (TITLE_A, ["Corpo 1.", "Corpo 2."]),
        )
        noticias = _parser_from_buf(buf)

        n = noticias[0]
        assert n.posicao_inicial < n.posicao_final

    def test_no_noticias_emits_warning(self, caplog: pytest.LogCaptureFixture) -> None:
        document = docx.Document()
        document.add_paragraph("Texto comum.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        with caplog.at_level(logging.WARNING):
            _parser_from_buf(buf)

        assert "Nenhuma notícia encontrada" in caplog.text

    def test_noticias_found_emits_info(self, caplog: pytest.LogCaptureFixture) -> None:
        buf = _build_docx((TITLE_A, ["Corpo."]))
        with caplog.at_level(logging.INFO):
            _parser_from_buf(buf)

        assert "extraída" in caplog.text


# ---------------------------------------------------------------------------
# Testes do NewsParser — capa do documento
# ---------------------------------------------------------------------------


class TestNewsParserCoverPage:
    """Testes de detecção e ignoramento da página de capa."""

    def test_content_before_bloco_marker_is_ignored(self) -> None:
        """Tudo antes de '1º BLOCO' deve ser descartado como capa."""
        document = docx.Document()
        # Capa: conteúdo que não deve virar notícia
        document.add_paragraph("BOLETIM DIA 10 DE JULHO de 2026")
        document.add_paragraph("DIÁLOGOS NOS TERRITÓRIOS: GESTÃO DO CUIDADO SAÚDE")
        document.add_paragraph("No dia 10 de julho, acontece o evento de saúde.")
        # Marcador de bloco
        document.add_paragraph("1º BLOCO – EVOLUÇÃO")
        # Notícia real
        document.add_paragraph(TITLE_A)
        document.add_paragraph("Corpo da notícia real.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert len(noticias) == 1
        assert noticias[0].titulo == TITLE_A

    def test_section_title_on_cover_is_not_detected_as_news(self) -> None:
        """Títulos de seção na capa (ex.: 'APRESENTAÇÃO DE LIBRAS') são descartados."""
        document = docx.Document()
        document.add_paragraph("APRESENTAÇÃO DE LIBRAS")  # 3 palavras – seção
        document.add_paragraph("1º BLOCO")
        document.add_paragraph(TITLE_A)
        document.add_paragraph("Corpo.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert len(noticias) == 1
        assert "APRESENTAÇÃO DE LIBRAS" not in noticias[0].titulo

    def test_bloco_marker_itself_is_not_in_body(self) -> None:
        """O texto do marcador de bloco não deve aparecer no corpo de nenhuma notícia."""
        document = docx.Document()
        document.add_paragraph("1º BLOCO – EVOLUÇÃO")
        document.add_paragraph(TITLE_A)
        document.add_paragraph("Corpo.")
        document.add_paragraph("2º BLOCO – SEGUNDO")
        document.add_paragraph(TITLE_B)
        document.add_paragraph("Corpo B.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        for n in noticias:
            assert "BLOCO" not in n.corpo

    def test_document_without_bloco_processes_from_start(self) -> None:
        """Sem marcador de bloco, o documento inteiro é processado."""
        buf = _build_docx(
            (TITLE_A, ["Corpo A"]),
            (TITLE_B, ["Corpo B"]),
        )
        noticias = _parser_from_buf(buf)

        assert len(noticias) == 2


# ---------------------------------------------------------------------------
# Testes do NewsParser — extração de tabelas
# ---------------------------------------------------------------------------


class TestNewsParserTableExtraction:
    """Testes do comportamento quando o documento contém tabelas."""

    def test_table_in_body_is_extracted_as_text(self) -> None:
        document = docx.Document()
        document.add_paragraph(TITLE_A)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cabeçalho A"
        table.cell(0, 1).text = "Cabeçalho B"
        table.cell(1, 0).text = "Valor 1"
        table.cell(1, 1).text = "Valor 2"
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert len(noticias) == 1
        assert "Cabeçalho A" in noticias[0].corpo
        assert "Valor 1" in noticias[0].corpo

    def test_table_cells_joined_by_tab_and_rows_by_newline(self) -> None:
        document = docx.Document()
        document.add_paragraph(TITLE_A)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert "A\tB" in noticias[0].corpo
        assert "C\tD" in noticias[0].corpo

    def test_empty_table_is_not_added_to_body(self) -> None:
        document = docx.Document()
        document.add_paragraph(TITLE_A)
        document.add_paragraph("Parágrafo real.")
        document.add_table(rows=1, cols=2)  # células vazias
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert noticias[0].corpo == "Parágrafo real."

    def test_table_before_first_title_is_ignored(self) -> None:
        document = docx.Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Dado órfão"
        document.add_paragraph(TITLE_A)
        document.add_paragraph("Corpo da notícia.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        assert "Dado órfão" not in noticias[0].corpo

    def test_table_followed_by_paragraph_in_same_news(self) -> None:
        document = docx.Document()
        document.add_paragraph(TITLE_A)
        document.add_paragraph("Introdução.")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Dado da tabela"
        document.add_paragraph("Conclusão.")
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        noticias = _parser_from_buf(buf)

        corpo = noticias[0].corpo
        assert "Introdução." in corpo
        assert "Dado da tabela" in corpo
        assert "Conclusão." in corpo

    def test_table_text_helper_with_populated_table(self) -> None:
        document = docx.Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "X"
        table.cell(0, 1).text = "Y"
        table.cell(1, 0).text = "Z"
        table.cell(1, 1).text = "W"

        result = _table_to_text(table)

        assert "X\tY" in result
        assert "Z\tW" in result

    def test_table_text_helper_with_empty_rows_skipped(self) -> None:
        document = docx.Document()
        table = document.add_table(rows=1, cols=2)
        # células em branco

        result = _table_to_text(table)

        assert result == ""


# ---------------------------------------------------------------------------
# Testes do NewsParser — integração com arquivo real no disco
# ---------------------------------------------------------------------------


class TestNewsParserWithDiskFile:
    """Testes que escrevem um .docx no sistema de arquivos temporário."""

    def test_parse_reads_file_from_disk(self, tmp_path: Path) -> None:
        buf = _build_docx(
            (TITLE_A, ["Conteúdo gravado no disco."]),
        )
        file_path = tmp_path / "teste.docx"
        file_path.write_bytes(buf.read())

        parser = NewsParser()
        noticias = parser.parse(file_path)

        assert len(noticias) == 1
        assert noticias[0].titulo == TITLE_A
        assert noticias[0].corpo == "Conteúdo gravado no disco."

    def test_parse_multiple_news_from_disk(self, tmp_path: Path) -> None:
        buf = _build_docx(
            (TITLE_A, ["Resultado dos jogos."]),
            (TITLE_B, ["Sessão no congresso."]),
            (TITLE_C, ["Índice subiu."]),
        )
        file_path = tmp_path / "boletim.docx"
        file_path.write_bytes(buf.read())

        parser = NewsParser()
        noticias = parser.parse(file_path)

        assert len(noticias) == 3
        assert noticias[1].titulo == TITLE_B


# ---------------------------------------------------------------------------
# Testes do NewsParser — arquivo BOLETIM real
# ---------------------------------------------------------------------------


@pytest.mark.skipif(
    not BOLETIM_PATH.exists(),
    reason=f"Arquivo BOLETIM não encontrado em {BOLETIM_PATH}",
)
class TestNewsParserBoletim:
    """Testes com o arquivo BOLETIM DIA 14 DE JULHO de 2026 – TERÇA-FEIRA.

    O parser deve retornar exatamente cinco notícias com os títulos,
    conteúdos e posições esperados.
    """

    @pytest.fixture()
    def noticias(self) -> list[Noticia]:
        parser = NewsParser()
        return parser.parse(BOLETIM_PATH)

    # --- Número de notícias ---

    def test_returns_exactly_five_news(self, noticias: list[Noticia]) -> None:
        assert len(noticias) == 5

    # --- Índices ordinais ---

    def test_indices_are_sequential_1_to_5(self, noticias: list[Noticia]) -> None:
        assert [n.indice for n in noticias] == [1, 2, 3, 4, 5]

    # --- Arquivo de origem ---

    def test_arquivo_origem_is_set(self, noticias: list[Noticia]) -> None:
        for n in noticias:
            assert n.arquivo_origem.endswith(".docx")

    # --- Notícia 1 ---

    def test_noticia1_title(self, noticias: list[Noticia]) -> None:
        assert noticias[0].titulo == (
            "PR: BANCO VAI TER QUE PAGAR DANOS MORAIS A CLIENTE PCD VISUAL"
        )

    def test_noticia1_corpo_contains_tribunal(self, noticias: list[Noticia]) -> None:
        assert "Câmara Cível" in noticias[0].corpo

    def test_noticia1_corpo_contains_lbi(self, noticias: list[Noticia]) -> None:
        assert "LBI" in noticias[0].corpo

    def test_noticia1_posicao_inicial(self, noticias: list[Noticia]) -> None:
        assert noticias[0].posicao_inicial == 19

    def test_noticia1_posicao_final(self, noticias: list[Noticia]) -> None:
        assert noticias[0].posicao_final == 25

    def test_noticia1_posicao_inicial_lt_posicao_final(
        self, noticias: list[Noticia]
    ) -> None:
        assert noticias[0].posicao_inicial < noticias[0].posicao_final

    # --- Notícia 2 ---

    def test_noticia2_title(self, noticias: list[Noticia]) -> None:
        assert noticias[1].titulo == (
            "PESSOAS COM DEFICIÊNCIA E A PREVENÇÃO DE DESASTRES: COMO FICA ?"
        )

    def test_noticia2_corpo_contains_postagem(self, noticias: list[Noticia]) -> None:
        assert "postagem no instagram" in noticias[1].corpo

    def test_noticia2_corpo_contains_protocolo(self, noticias: list[Noticia]) -> None:
        assert "protocolo" in noticias[1].corpo

    def test_noticia2_posicao_inicial(self, noticias: list[Noticia]) -> None:
        assert noticias[1].posicao_inicial == 28

    def test_noticia2_posicao_final(self, noticias: list[Noticia]) -> None:
        assert noticias[1].posicao_final == 33

    # --- Notícia 3 ---

    def test_noticia3_title(self, noticias: list[Noticia]) -> None:
        assert noticias[2].titulo == (
            "PEDREIRO RECEBE PRÓTESE DE R$ 80 MIL DOADA POR LUCIANO HANG DA HAVAN"
        )

    def test_noticia3_corpo_contains_pedreiro(self, noticias: list[Noticia]) -> None:
        assert "Douglas" in noticias[2].corpo

    def test_noticia3_corpo_contains_havan(self, noticias: list[Noticia]) -> None:
        assert "HAVAN" in noticias[2].corpo

    def test_noticia3_posicao_inicial(self, noticias: list[Noticia]) -> None:
        assert noticias[2].posicao_inicial == 37

    def test_noticia3_posicao_final(self, noticias: list[Noticia]) -> None:
        assert noticias[2].posicao_final == 47

    # --- Notícia 4 ---

    def test_noticia4_title(self, noticias: list[Noticia]) -> None:
        assert noticias[3].titulo == (
            "APROVADA REGULAMENTAÇÃO DA PROFISSÃO DE AUDIODESCRITOR"
        )

    def test_noticia4_corpo_contains_senado(self, noticias: list[Noticia]) -> None:
        assert "Senado" in noticias[3].corpo

    def test_noticia4_corpo_contains_cdh(self, noticias: list[Noticia]) -> None:
        assert "CDH" in noticias[3].corpo

    def test_noticia4_posicao_inicial(self, noticias: list[Noticia]) -> None:
        assert noticias[3].posicao_inicial == 49

    def test_noticia4_posicao_final(self, noticias: list[Noticia]) -> None:
        assert noticias[3].posicao_final == 63

    # --- Notícia 5 ---

    def test_noticia5_title(self, noticias: list[Noticia]) -> None:
        assert noticias[4].titulo == (
            "CONGRESSO SÃO PAULO TEAMA ACONTECE NA CAPITAL PAULISTA"
        )

    def test_noticia5_corpo_contains_autismo(self, noticias: list[Noticia]) -> None:
        assert "autismo" in noticias[4].corpo

    def test_noticia5_corpo_contains_sympla(self, noticias: list[Noticia]) -> None:
        assert "sympla" in noticias[4].corpo

    def test_noticia5_posicao_inicial(self, noticias: list[Noticia]) -> None:
        assert noticias[4].posicao_inicial == 67

    def test_noticia5_posicao_final(self, noticias: list[Noticia]) -> None:
        assert noticias[4].posicao_final == 75

    # --- Ordenação e integridade global ---

    def test_all_titles_are_all_caps(self, noticias: list[Noticia]) -> None:
        for n in noticias:
            letters = [c for c in n.titulo if c.isalpha()]
            assert all(c.isupper() for c in letters), (
                f"Título não está em caixa alta: {n.titulo!r}"
            )

    def test_positions_increase_monotonically(self, noticias: list[Noticia]) -> None:
        for i in range(len(noticias) - 1):
            assert noticias[i].posicao_final < noticias[i + 1].posicao_inicial, (
                f"Posições sobrepostas entre notícia {i + 1} e {i + 2}"
            )

    def test_all_news_have_non_empty_corpo(self, noticias: list[Noticia]) -> None:
        for n in noticias:
            assert n.corpo.strip(), f"Corpo vazio na notícia: {n.titulo!r}"

    def test_cover_page_content_not_in_any_news(self, noticias: list[Noticia]) -> None:
        """Conteúdo da capa não deve aparecer nos títulos das notícias."""
        titles = [n.titulo for n in noticias]
        assert not any("BOLETIM" in t for t in titles)
        assert not any("APRESENTAÇÃO DE LIBRAS" == t for t in titles)
        assert not any("DIÁLOGOS NOS TERRITÓRIOS" in t for t in titles)

    def test_bloco_markers_not_in_any_corpo(self, noticias: list[Noticia]) -> None:
        for n in noticias:
            assert "1º BLOCO" not in n.corpo
            assert "2º BLOCO" not in n.corpo
            assert "3º BLOCO" not in n.corpo
