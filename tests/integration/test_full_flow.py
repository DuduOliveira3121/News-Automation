"""Testes de integração do fluxo completo de automação de notícias.

Fluxo exercitado
================
    .docx → ParseDocxUseCase → DB
         → ReviewNewsUseCase (IA mockada) → NewsService.update_fields → DB
         → PublishNewsUseCase (SitePublisher mockado) → DB

Princípios de isolamento
========================
* **Banco de dados** — SQLite em memória (sem arquivo em disco).
* **IA (OpenAI)** — IAIClient substituído por MagicMock com respostas fixas.
* **Browser (Playwright)** — SitePublisher substituído por MagicMock; nenhum
  processo de browser é aberto.
* **Arquivo .docx** — criado em tempo de execução via ``python-docx`` em um
  diretório temporário gerenciado pelo pytest (``tmp_path``).
* **Settings** — portal_url e openai_api_key patchados onde necessário.

Cobertura de cenários
=====================
+--------------------------------------------------+------------------------------------------+
| Teste                                            | Camadas exercitadas                      |
+--------------------------------------------------+------------------------------------------+
| Parsing extrai notícias corretas do .docx        | Parser → DocxService → ParseDocxUseCase  |
| Múltiplos Heading 1 geram múltiplas entidades    | Parser → DocxService → ParseDocxUseCase  |
| Notícias são persistidas no banco                | ParseDocxUseCase → NewsRepository        |
| Revisão IA retorna DTO com todos os campos       | AIReviewService → ReviewNewsUseCase      |
| Campos gerados pela IA são salvos no banco       | NewsService → NewsRepository             |
| Publicação chama SitePublisher com dados certos  | PublicationService → SitePublisher (mock)|
| Publicação bem-sucedida muda status para PUBLISHED| PublishNewsUseCase → NewsRepository      |
| Falha na publicação muda status para FAILED      | PublishNewsUseCase → NewsRepository      |
| Registro de Publication é criado no banco        | PublishNewsUseCase → PublicationRepository|
| Fluxo completo ponta-a-ponta sem serviços reais  | Todas as camadas                         |
+--------------------------------------------------+------------------------------------------+

Execução
--------
    pytest tests/integration/test_full_flow.py -v
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Generator
from unittest.mock import MagicMock, call, patch, patch as _patch

import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker

# ─── Configuração ─────────────────────────────────────────────────────────────
from config.settings import settings as _app_settings

# ─── Domínio ──────────────────────────────────────────────────────────────────
from src.domain.entities.news import News, NewsStatus
from src.domain.entities.publication import PublicationStatus

# ─── Aplicação ────────────────────────────────────────────────────────────────
from src.application.dtos.news_dto import NewsDTO
from src.application.services.ai_review_service import AIReviewService
from src.application.services.docx_service import DocxService
from src.application.services.news_service import NewsService
from src.application.services.publication_service import PublicationService
from src.application.use_cases.parse_docx_use_case import ParseDocxUseCase
from src.application.use_cases.publish_news_use_case import PublishNewsUseCase
from src.application.use_cases.review_news_use_case import ReviewNewsUseCase

# ─── Infraestrutura ───────────────────────────────────────────────────────────
from src.infrastructure.database.models import Base
from src.infrastructure.repositories.news_repository import NewsRepository
from src.infrastructure.repositories.publication_repository import (
    SQLAlchemyPublicationRepository,
)

# ──────────────────────────────────────────────────────────────────────────────
# Constantes de teste
# ──────────────────────────────────────────────────────────────────────────────

_PORTAL_URL = "https://portal.test.local"

_AI_TITULO = "Título Gerado pela IA"
_AI_RESUMO = "Resumo gerado pela IA em duas frases."
_AI_CONTEUDO = "Conteúdo melhorado pela IA sem erros gramaticais."
_AI_ALT_TEXT = "Imagem relacionada à notícia gerada pela IA."


# ──────────────────────────────────────────────────────────────────────────────
# Fixtures compartilhadas
# ──────────────────────────────────────────────────────────────────────────────


@pytest.fixture()
def db_session() -> Generator[Session, None, None]:
    """Sessão SQLAlchemy com SQLite em memória — banco zerado a cada teste."""
    engine = create_engine(
        "sqlite:///:memory:", connect_args={"check_same_thread": False}
    )
    Base.metadata.create_all(engine)
    factory = sessionmaker(bind=engine, autocommit=False, autoflush=False)
    session = factory()
    try:
        yield session
    finally:
        session.close()
        Base.metadata.drop_all(engine)
        engine.dispose()


@pytest.fixture()
def news_repo(db_session: Session) -> NewsRepository:
    return NewsRepository(db_session)


@pytest.fixture()
def pub_repo(db_session: Session) -> SQLAlchemyPublicationRepository:
    return SQLAlchemyPublicationRepository(db_session)


@pytest.fixture()
def news_service(news_repo: NewsRepository) -> NewsService:
    return NewsService(news_repo)


@pytest.fixture()
def mock_ai_client() -> MagicMock:
    """IAIClient substituto com respostas fixas e determinísticas.

    O método ``complete(system_prompt, user_message)`` é roteado por
    palavras-chave nos prompts de sistema para retornar valores distintos
    por tipo de geração (título, resumo, conteúdo, alt-text).
    """
    client = MagicMock()

    def _complete_side_effect(system_prompt: str, user_message: str) -> str:
        prompt_lower = system_prompt.lower()
        if "título" in prompt_lower or "titulo" in prompt_lower:
            return _AI_TITULO
        if "resumo" in prompt_lower:
            return _AI_RESUMO
        if "melhore" in prompt_lower or "melhor" in prompt_lower:
            return _AI_CONTEUDO
        if "alternativo" in prompt_lower or "alt text" in prompt_lower:
            return _AI_ALT_TEXT
        return "Resposta genérica da IA."

    client.complete.side_effect = _complete_side_effect
    return client


@pytest.fixture()
def ai_review_service(mock_ai_client: MagicMock) -> AIReviewService:
    """AIReviewService com cliente de IA mockado."""
    return AIReviewService(mock_ai_client)


@pytest.fixture()
def mock_site_publisher_cls() -> MagicMock:
    """Substituto para a classe SitePublisher que simula publicação bem-sucedida.

    Configura o context manager (__enter__/__exit__) e todos os métodos
    chamados pelo PublicationService.publish().
    """
    publisher_instance = MagicMock()
    publisher_instance.__enter__ = MagicMock(return_value=publisher_instance)
    publisher_instance.__exit__ = MagicMock(return_value=False)

    publisher_cls = MagicMock(return_value=publisher_instance)
    return publisher_cls


@pytest.fixture()
def patched_portal(mock_site_publisher_cls: MagicMock):
    """Fixture que aplica todos os patches necessários para publicação.

    - Patcha ``settings.portal_url / username / password`` via patch.object
      (modifica o singleton em todos os módulos que o importaram).
    - Patcha ``SitePublisher`` para evitar abertura de browser real.

    Uso::

        def test_algo(patched_portal):
            with patched_portal:
                publish_uc.execute(news_id)
    """
    class _PatchedPortal:
        def __enter__(self) -> "_PatchedPortal":
            self._p_publisher = patch(
                "src.infrastructure.automation.site_publisher.SitePublisher",
                mock_site_publisher_cls,
            )
            self._p_url = patch.object(_app_settings, "portal_url", _PORTAL_URL)
            self._p_user = patch.object(_app_settings, "portal_username", "editor")
            self._p_pass = patch.object(_app_settings, "portal_password", "senha123")

            self._p_publisher.start()
            self._p_url.start()
            self._p_user.start()
            self._p_pass.start()
            return self

        def __exit__(self, *args: object) -> None:
            self._p_publisher.stop()
            self._p_url.stop()
            self._p_user.stop()
            self._p_pass.stop()

    return _PatchedPortal()


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Cria um arquivo .docx real com uma única notícia (Heading 1 + corpo)."""
    import docx

    doc = docx.Document()
    doc.add_heading("Prefeito anuncia novo programa habitacional", level=1)
    doc.add_paragraph(
        "O prefeito municipal anunciou nesta terça-feira a criação de um novo "
        "programa habitacional que prevê a construção de 500 unidades."
    )
    doc.add_paragraph(
        "O programa será financiado com recursos federais e municipais, "
        "com previsão de entrega das primeiras unidades em 18 meses."
    )

    path = tmp_path / "boletim_simples.docx"
    doc.save(str(path))
    return path


@pytest.fixture()
def multi_news_docx(tmp_path: Path) -> Path:
    """Cria um .docx com três notícias separadas por Heading 1."""
    import docx

    doc = docx.Document()

    doc.add_heading("Câmara aprova orçamento municipal", level=1)
    doc.add_paragraph("A câmara municipal aprovou ontem o orçamento para o próximo exercício.")
    doc.add_paragraph("O valor total aprovado é de R$ 850 milhões.")

    doc.add_heading("Obras na Avenida Central iniciam segunda-feira", level=1)
    doc.add_paragraph("A prefeitura iniciará obras de recapeamento na Avenida Central na próxima semana.")

    doc.add_heading("Festival Cultural reúne 10 mil pessoas no fim de semana", level=1)
    doc.add_paragraph("O festival cultural realizado no parque central reuniu cerca de 10 mil visitantes.")
    doc.add_paragraph("A edição deste ano contou com shows, exposições e gastronomia local.")

    path = tmp_path / "boletim_multiplo.docx"
    doc.save(str(path))
    return path


# ──────────────────────────────────────────────────────────────────────────────
# 1. Parsing do Word e persistência no banco
# ──────────────────────────────────────────────────────────────────────────────


class TestDocxParsingAndPersistence:
    """Testes da etapa Word → Extração → Banco de dados."""

    def test_parse_extrai_titulo_e_conteudo(
        self,
        sample_docx: Path,
        news_repo: NewsRepository,
    ) -> None:
        """.docx com uma notícia deve gerar uma entidade com título correto."""
        uc = ParseDocxUseCase(DocxService(), news_repo)

        saved = uc.execute(sample_docx)

        assert len(saved) == 1
        assert saved[0].titulo == "Prefeito anuncia novo programa habitacional"
        assert "programa habitacional" in saved[0].conteudo

    def test_parse_preserva_nome_do_arquivo_de_origem(
        self,
        sample_docx: Path,
        news_repo: NewsRepository,
    ) -> None:
        """source_file deve conter o nome do .docx processado."""
        uc = ParseDocxUseCase(DocxService(), news_repo)

        saved = uc.execute(sample_docx)

        assert saved[0].source_file == "boletim_simples.docx"

    def test_parse_status_inicial_e_pending(
        self,
        sample_docx: Path,
        news_repo: NewsRepository,
    ) -> None:
        """Notícia recém-extraída deve ter status PENDING."""
        uc = ParseDocxUseCase(DocxService(), news_repo)

        saved = uc.execute(sample_docx)

        assert saved[0].status == NewsStatus.PENDING

    def test_parse_multiplas_noticias_cria_multiplas_entidades(
        self,
        multi_news_docx: Path,
        news_repo: NewsRepository,
    ) -> None:
        """.docx com três Heading 1 deve gerar três entidades distintas."""
        uc = ParseDocxUseCase(DocxService(), news_repo)

        saved = uc.execute(multi_news_docx)

        assert len(saved) == 3
        titulos = [n.titulo for n in saved]
        assert "Câmara aprova orçamento municipal" in titulos
        assert "Obras na Avenida Central iniciam segunda-feira" in titulos
        assert "Festival Cultural reúne 10 mil pessoas no fim de semana" in titulos

    def test_noticias_persistidas_sao_recuperaveis_do_banco(
        self,
        multi_news_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
    ) -> None:
        """Após commit, todas as notícias devem ser recuperáveis via find_all."""
        uc = ParseDocxUseCase(DocxService(), news_repo)

        saved = uc.execute(multi_news_docx)
        db_session.commit()

        all_news = news_repo.find_all()
        assert len(all_news) == 3

        ids_saved = {n.id for n in saved}
        ids_found = {n.id for n in all_news}
        assert ids_saved == ids_found

    def test_cada_noticia_tem_id_unico(
        self,
        multi_news_docx: Path,
        news_repo: NewsRepository,
    ) -> None:
        """IDs gerados automaticamente devem ser únicos entre as entidades."""
        uc = ParseDocxUseCase(DocxService(), news_repo)

        saved = uc.execute(multi_news_docx)

        ids = [n.id for n in saved]
        assert len(ids) == len(set(ids))

    def test_parse_arquivo_invalido_levanta_value_error(
        self,
        tmp_path: Path,
        news_repo: NewsRepository,
    ) -> None:
        """Arquivo com extensão inválida deve levantar ValueError."""
        invalid_file = tmp_path / "doc.txt"
        invalid_file.write_text("conteúdo inválido")

        uc = ParseDocxUseCase(DocxService(), news_repo)

        with pytest.raises(ValueError, match="Extensão inválida"):
            uc.execute(invalid_file)

    def test_parse_arquivo_inexistente_levanta_file_not_found(
        self,
        tmp_path: Path,
        news_repo: NewsRepository,
    ) -> None:
        """Caminho inexistente deve levantar FileNotFoundError."""
        uc = ParseDocxUseCase(DocxService(), news_repo)

        with pytest.raises(FileNotFoundError):
            uc.execute(tmp_path / "nao_existe.docx")


# ──────────────────────────────────────────────────────────────────────────────
# 2. Geração de conteúdo via IA (cliente mockado)
# ──────────────────────────────────────────────────────────────────────────────


class TestAIReviewIntegration:
    """Testes da etapa de revisão/geração de conteúdo via IA."""

    def _save_one_news(
        self,
        sample_docx: Path,
        news_repo: NewsRepository,
        db_session: Session,
    ) -> News:
        """Auxiliar: faz o parse e retorna a primeira notícia salva."""
        uc = ParseDocxUseCase(DocxService(), news_repo)
        saved = uc.execute(sample_docx)
        db_session.commit()
        return saved[0]

    def test_review_uc_retorna_dto_com_titulo_gerado(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        ai_review_service: AIReviewService,
    ) -> None:
        """ReviewNewsUseCase deve retornar ReviewResultDTO com título não vazio."""
        news = self._save_one_news(sample_docx, news_repo, db_session)
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)

        result = review_uc.execute(news.id)

        assert result.titulo == _AI_TITULO

    def test_review_uc_retorna_dto_com_resumo_gerado(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        ai_review_service: AIReviewService,
    ) -> None:
        """ReviewResultDTO deve conter o resumo gerado pela IA."""
        news = self._save_one_news(sample_docx, news_repo, db_session)
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)

        result = review_uc.execute(news.id)

        assert result.resumo == _AI_RESUMO

    def test_review_uc_retorna_dto_com_conteudo_melhorado(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        ai_review_service: AIReviewService,
    ) -> None:
        """ReviewResultDTO deve conter o conteúdo revisado pela IA."""
        news = self._save_one_news(sample_docx, news_repo, db_session)
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)

        result = review_uc.execute(news.id)

        assert result.reviewed_content == _AI_CONTEUDO

    def test_review_uc_retorna_dto_com_texto_alternativo(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        ai_review_service: AIReviewService,
    ) -> None:
        """ReviewResultDTO deve conter o alt-text gerado pela IA."""
        news = self._save_one_news(sample_docx, news_repo, db_session)
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)

        result = review_uc.execute(news.id)

        assert result.texto_alternativo == _AI_ALT_TEXT

    def test_review_uc_nao_persiste_automaticamente(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        ai_review_service: AIReviewService,
    ) -> None:
        """Chamar review_uc.execute() não deve alterar os dados no banco.

        O resultado é retornado para o editor; a persistência só ocorre
        quando o usuário clica em 'Salvar'.
        """
        news = self._save_one_news(sample_docx, news_repo, db_session)
        original_titulo = news.titulo
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)

        review_uc.execute(news.id)

        news_no_db = news_repo.find_by_id(news.id)
        assert news_no_db.titulo == original_titulo

    def test_news_service_persiste_campos_revisados(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        news_service: NewsService,
        ai_review_service: AIReviewService,
    ) -> None:
        """Após update_fields, os dados da IA devem ser recuperáveis do banco."""
        uc_parse = ParseDocxUseCase(DocxService(), news_repo)
        saved = uc_parse.execute(sample_docx)
        db_session.commit()
        news = saved[0]

        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)
        result = review_uc.execute(news.id)

        # Simula o clique em "Salvar" na UI
        news_service.update_fields(
            news_id=news.id,
            titulo=result.titulo,
            conteudo=result.reviewed_content,
            resumo=result.resumo,
            texto_alternativo=result.texto_alternativo,
        )
        db_session.commit()

        updated = news_repo.find_by_id(news.id)
        assert updated.titulo == _AI_TITULO
        assert updated.conteudo == _AI_CONTEUDO
        assert updated.resumo == _AI_RESUMO
        assert updated.texto_alternativo == _AI_ALT_TEXT

    def test_review_uc_levanta_erro_para_id_inexistente(
        self,
        db_session: Session,
        news_repo: NewsRepository,
        ai_review_service: AIReviewService,
    ) -> None:
        """ReviewNewsUseCase deve levantar ValueError para ID desconhecido."""
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)

        with pytest.raises(ValueError, match="não encontrada"):
            review_uc.execute("id-que-nao-existe")

    def test_ai_client_recebe_conteudo_da_noticia(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        mock_ai_client: MagicMock,
    ) -> None:
        """O conteúdo original da notícia deve ser passado ao cliente de IA."""
        uc_parse = ParseDocxUseCase(DocxService(), news_repo)
        saved = uc_parse.execute(sample_docx)
        db_session.commit()
        news = saved[0]

        ai_review_service = AIReviewService(mock_ai_client)
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)
        review_uc.execute(news.id)

        # Cada chamada ao complete() deve receber o conteúdo original
        calls = mock_ai_client.complete.call_args_list
        assert len(calls) == 4  # titulo, resumo, conteúdo melhorado, alt-text
        for c in calls:
            _, kwargs = c
            user_msg = kwargs.get("user_message") or c.args[1]
            assert news.conteudo.strip()[:20] in user_msg


# ──────────────────────────────────────────────────────────────────────────────
# 3. Publicação com SitePublisher mockado
# ──────────────────────────────────────────────────────────────────────────────


class TestPublicationIntegration:
    """Testes da etapa de publicação no portal (SitePublisher mockado)."""

    def _prepare_news(
        self,
        sample_docx: Path,
        news_repo: NewsRepository,
        news_service: NewsService,
        ai_review_service: AIReviewService,
        db_session: Session,
    ) -> News:
        """Auxiliar: parse → salva → revisa → update_fields → retorna entidade."""
        uc_parse = ParseDocxUseCase(DocxService(), news_repo)
        saved = uc_parse.execute(sample_docx)
        db_session.commit()
        news = saved[0]

        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)
        result = review_uc.execute(news.id)
        news_service.update_fields(
            news_id=news.id,
            titulo=result.titulo,
            conteudo=result.reviewed_content,
            resumo=result.resumo,
            texto_alternativo=result.texto_alternativo,
        )
        db_session.commit()
        return news_repo.find_by_id(news.id)

    def test_publicacao_bem_sucedida_muda_status_para_published(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        news_service: NewsService,
        pub_repo: SQLAlchemyPublicationRepository,
        ai_review_service: AIReviewService,
        mock_site_publisher_cls: MagicMock,
    ) -> None:
        """Publicação bem-sucedida deve transicionar a notícia para PUBLISHED."""
        news = self._prepare_news(
            sample_docx, news_repo, news_service, ai_review_service, db_session
        )
        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with (
            patch(
                "src.infrastructure.automation.site_publisher.SitePublisher",
                mock_site_publisher_cls,
            ),
            patch(
                "src.application.services.publication_service.settings"
            ) as mock_settings,
        ):
            mock_settings.portal_url = _PORTAL_URL
            mock_settings.portal_username = "editor"
            mock_settings.portal_password = "senha123"

            result = publish_uc.execute(news.id)
            db_session.commit()

        news_after = news_repo.find_by_id(news.id)
        assert news_after.status == NewsStatus.PUBLISHED
        assert result.status == PublicationStatus.SUCCESS

    def test_publicacao_cria_registro_no_banco(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        news_service: NewsService,
        pub_repo: SQLAlchemyPublicationRepository,
        ai_review_service: AIReviewService,
        mock_site_publisher_cls: MagicMock,
    ) -> None:
        """Deve existir exatamente um registro de Publication no banco após publicar."""
        news = self._prepare_news(
            sample_docx, news_repo, news_service, ai_review_service, db_session
        )
        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with (
            patch(
                "src.infrastructure.automation.site_publisher.SitePublisher",
                mock_site_publisher_cls,
            ),
            patch(
                "src.application.services.publication_service.settings"
            ) as mock_settings,
        ):
            mock_settings.portal_url = _PORTAL_URL
            mock_settings.portal_username = "editor"
            mock_settings.portal_password = "senha123"

            result = publish_uc.execute(news.id)
            db_session.commit()

        publications = pub_repo.find_by_news_id(news.id)
        assert len(publications) == 1
        assert publications[0].news_id == news.id
        assert publications[0].portal_url == _PORTAL_URL

    def test_publicacao_chama_site_publisher_com_dados_corretos(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        news_service: NewsService,
        pub_repo: SQLAlchemyPublicationRepository,
        ai_review_service: AIReviewService,
        mock_site_publisher_cls: MagicMock,
    ) -> None:
        """SitePublisher deve receber o título e conteúdo atualizados pela IA."""
        news = self._prepare_news(
            sample_docx, news_repo, news_service, ai_review_service, db_session
        )
        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with (
            patch(
                "src.infrastructure.automation.site_publisher.SitePublisher",
                mock_site_publisher_cls,
            ),
            patch(
                "src.application.services.publication_service.settings"
            ) as mock_settings,
        ):
            mock_settings.portal_url = _PORTAL_URL
            mock_settings.portal_username = "editor"
            mock_settings.portal_password = "senha123"

            publish_uc.execute(news.id)

        publisher_instance = mock_site_publisher_cls.return_value
        publisher_instance.preencher_campos.assert_called_once()
        dto_passado: NewsDTO = publisher_instance.preencher_campos.call_args[0][0]
        assert dto_passado.titulo == _AI_TITULO
        assert dto_passado.conteudo == _AI_CONTEUDO

    def test_publicacao_chama_sequencia_correta_de_metodos(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        news_service: NewsService,
        pub_repo: SQLAlchemyPublicationRepository,
        ai_review_service: AIReviewService,
        mock_site_publisher_cls: MagicMock,
    ) -> None:
        """login → abrir_painel → criar_noticia → preencher_campos → publicar."""
        news = self._prepare_news(
            sample_docx, news_repo, news_service, ai_review_service, db_session
        )
        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with (
            patch(
                "src.infrastructure.automation.site_publisher.SitePublisher",
                mock_site_publisher_cls,
            ),
            patch(
                "src.application.services.publication_service.settings"
            ) as mock_settings,
        ):
            mock_settings.portal_url = _PORTAL_URL
            mock_settings.portal_username = "editor"
            mock_settings.portal_password = "senha123"

            publish_uc.execute(news.id)

        pub_instance = mock_site_publisher_cls.return_value
        method_calls = [c[0] for c in pub_instance.method_calls if not c[0].startswith("__")]
        expected_order = ["login", "abrir_painel", "criar_noticia", "preencher_campos", "publicar"]
        assert method_calls == expected_order

    def test_falha_na_publicacao_muda_status_para_failed(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        news_service: NewsService,
        pub_repo: SQLAlchemyPublicationRepository,
        ai_review_service: AIReviewService,
        mock_site_publisher_cls: MagicMock,
    ) -> None:
        """RuntimeError no SitePublisher deve marcar a notícia como FAILED."""
        news = self._prepare_news(
            sample_docx, news_repo, news_service, ai_review_service, db_session
        )

        # Configura o publisher para lançar exceção ao publicar
        failing_instance = mock_site_publisher_cls.return_value
        failing_instance.publicar.side_effect = RuntimeError("Timeout ao clicar em Publicar")

        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with (
            patch(
                "src.infrastructure.automation.site_publisher.SitePublisher",
                mock_site_publisher_cls,
            ),
            patch(
                "src.application.services.publication_service.settings"
            ) as mock_settings,
        ):
            mock_settings.portal_url = _PORTAL_URL
            mock_settings.portal_username = "editor"
            mock_settings.portal_password = "senha123"

            result = publish_uc.execute(news.id)
            db_session.commit()

        news_after = news_repo.find_by_id(news.id)
        assert news_after.status == NewsStatus.FAILED
        assert result.status == PublicationStatus.FAILED
        assert "Timeout" in result.error_message

    def test_publicacao_sem_portal_url_levanta_value_error(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        news_service: NewsService,
        pub_repo: SQLAlchemyPublicationRepository,
        ai_review_service: AIReviewService,
    ) -> None:
        """PublicationService deve rejeitar a publicação se PORTAL_URL não estiver configurada."""
        news = self._prepare_news(
            sample_docx, news_repo, news_service, ai_review_service, db_session
        )
        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with patch(
            "src.application.services.publication_service.settings"
        ) as mock_settings:
            mock_settings.portal_url = ""

            with pytest.raises(ValueError, match="PORTAL_URL"):
                publish_uc.execute(news.id)

    def test_publicacao_noticia_inexistente_levanta_value_error(
        self,
        db_session: Session,
        news_repo: NewsRepository,
        pub_repo: SQLAlchemyPublicationRepository,
    ) -> None:
        """PublishNewsUseCase deve levantar ValueError para news_id desconhecido."""
        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with pytest.raises(ValueError, match="não encontrada"):
            publish_uc.execute("id-inexistente")


# ──────────────────────────────────────────────────────────────────────────────
# 4. Fluxo completo ponta-a-ponta
# ──────────────────────────────────────────────────────────────────────────────


class TestCompleteFlow:
    """Teste ponta-a-ponta: .docx → IA → banco → SitePublisher (tudo mockado)."""

    def test_fluxo_completo_sem_servicos_externos(
        self,
        sample_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        pub_repo: SQLAlchemyPublicationRepository,
        news_service: NewsService,
        ai_review_service: AIReviewService,
        mock_site_publisher_cls: MagicMock,
    ) -> None:
        """Executa o pipeline completo verificando cada etapa do fluxo.

        Etapas:
            1. Lê .docx real e persiste no banco (ParseDocxUseCase)
            2. Verifica status inicial PENDING
            3. Gera conteúdo via IA mockada (ReviewNewsUseCase)
            4. Salva campos gerados no banco (NewsService.update_fields)
            5. Publica com SitePublisher mockado (PublishNewsUseCase)
            6. Verifica status final PUBLISHED e registro de Publication
        """
        # ── Etapa 1: Parse do Word ────────────────────────────────────────────
        parse_uc = ParseDocxUseCase(DocxService(), news_repo)
        saved = parse_uc.execute(sample_docx)
        db_session.commit()

        assert len(saved) == 1, "Deve haver exatamente uma notícia extraída"
        news = saved[0]
        news_id = news.id

        # ── Etapa 2: Verificação do estado inicial ────────────────────────────
        news_initial = news_repo.find_by_id(news_id)
        assert news_initial.status == NewsStatus.PENDING
        assert news_initial.titulo == "Prefeito anuncia novo programa habitacional"

        # ── Etapa 3: Geração de conteúdo via IA ──────────────────────────────
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)
        review_result = review_uc.execute(news_id)

        assert review_result.titulo == _AI_TITULO
        assert review_result.resumo == _AI_RESUMO
        assert review_result.reviewed_content == _AI_CONTEUDO
        assert review_result.texto_alternativo == _AI_ALT_TEXT

        # Banco ainda deve ter o título original (sem persistência automática)
        assert news_repo.find_by_id(news_id).titulo == "Prefeito anuncia novo programa habitacional"

        # ── Etapa 4: Salvar campos revisados ──────────────────────────────────
        news_service.update_fields(
            news_id=news_id,
            titulo=review_result.titulo,
            conteudo=review_result.reviewed_content,
            resumo=review_result.resumo,
            texto_alternativo=review_result.texto_alternativo,
        )
        db_session.commit()

        news_saved = news_repo.find_by_id(news_id)
        assert news_saved.titulo == _AI_TITULO
        assert news_saved.conteudo == _AI_CONTEUDO
        assert news_saved.resumo == _AI_RESUMO

        # ── Etapa 5: Publicação no portal ─────────────────────────────────────
        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with (
            patch(
                "src.infrastructure.automation.site_publisher.SitePublisher",
                mock_site_publisher_cls,
            ),
            patch(
                "src.application.services.publication_service.settings"
            ) as mock_settings,
            patch(
                "src.application.use_cases.publish_news_use_case.settings"
            ) as mock_uc_settings,
        ):
            mock_settings.portal_url = _PORTAL_URL
            mock_settings.portal_username = "editor"
            mock_settings.portal_password = "senha123"
            mock_uc_settings.portal_url = _PORTAL_URL

            publication = publish_uc.execute(news_id)
            db_session.commit()

        # ── Etapa 6: Verificação do estado final ──────────────────────────────
        news_final = news_repo.find_by_id(news_id)
        assert news_final.status == NewsStatus.PUBLISHED, (
            f"Esperado PUBLISHED, obtido {news_final.status}"
        )

        assert publication.status == PublicationStatus.SUCCESS
        assert publication.news_id == news_id
        assert publication.portal_url == _PORTAL_URL
        assert publication.published_at is not None

        # Publication deve estar registrada no banco
        db_publications = pub_repo.find_by_news_id(news_id)
        assert len(db_publications) == 1
        assert db_publications[0].status == PublicationStatus.SUCCESS

    def test_fluxo_com_multiplas_noticias_publica_apenas_a_selecionada(
        self,
        multi_news_docx: Path,
        db_session: Session,
        news_repo: NewsRepository,
        pub_repo: SQLAlchemyPublicationRepository,
        news_service: NewsService,
        ai_review_service: AIReviewService,
        mock_site_publisher_cls: MagicMock,
    ) -> None:
        """Ao publicar uma notícia de um arquivo com três, só ela muda de status."""
        parse_uc = ParseDocxUseCase(DocxService(), news_repo)
        all_saved = parse_uc.execute(multi_news_docx)
        db_session.commit()

        assert len(all_saved) == 3
        target = all_saved[1]  # Seleciona a segunda notícia

        # Revisa e salva apenas a notícia alvo
        review_uc = ReviewNewsUseCase(ai_review_service, news_repo)
        result = review_uc.execute(target.id)
        news_service.update_fields(
            news_id=target.id,
            titulo=result.titulo,
            conteudo=result.reviewed_content,
        )
        db_session.commit()

        # Publica apenas a notícia alvo
        pub_service = PublicationService()
        publish_uc = PublishNewsUseCase(pub_service, news_repo, pub_repo)

        with (
            patch(
                "src.infrastructure.automation.site_publisher.SitePublisher",
                mock_site_publisher_cls,
            ),
            patch(
                "src.application.services.publication_service.settings"
            ) as mock_settings,
            patch(
                "src.application.use_cases.publish_news_use_case.settings"
            ) as mock_uc_settings,
        ):
            mock_settings.portal_url = _PORTAL_URL
            mock_settings.portal_username = "editor"
            mock_settings.portal_password = "senha123"
            mock_uc_settings.portal_url = _PORTAL_URL

            publish_uc.execute(target.id)
            db_session.commit()

        # Apenas a notícia publicada deve estar com status PUBLISHED
        for news in all_saved:
            n = news_repo.find_by_id(news.id)
            if news.id == target.id:
                assert n.status == NewsStatus.PUBLISHED
            else:
                assert n.status == NewsStatus.PENDING, (
                    f"Notícia {n.titulo!r} não deveria ter sido alterada"
                )
